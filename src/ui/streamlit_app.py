"""
Streamlit web interface for the Research Paper Summarizer.

This module provides a user-friendly web interface for uploading PDFs,
running analysis tasks, and viewing results with live preview capabilities.
"""

import streamlit as st
import tempfile
import base64
import json
import io
from pathlib import Path
import sys
from typing import Dict, Any, Optional
import time
import traceback

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from agent.input_processor import process_paper
from agent.planner import plan_analysis
from agent.executor import create_executor
from agent.run_agent import run_analysis

# Document generation
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Configure Streamlit page
st.set_page_config(
    page_title="Research Paper Summarizer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    .task-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: #2c3e50;
        margin: 1rem 0;
        border-bottom: 2px solid #3498db;
        padding-bottom: 0.5rem;
    }
    
    .metric-card {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #3498db;
        margin: 0.5rem 0;
    }
    
    .pdf-viewer {
        border: 1px solid #ddd;
        border-radius: 0.5rem;
        height: 600px;
        overflow: auto;
    }
    
    .stAlert > div {
        padding: 1rem;
    }
</style>
""", unsafe_allow_html=True)


class StreamlitApp:
    """Main Streamlit application class."""
    
    def __init__(self):
        self.init_session_state()
    
    def init_session_state(self):
        """Initialize Streamlit session state variables."""
        defaults = {
            'analysis_results': None,
            'uploaded_file': None,
            'model_loaded': False,
            'current_model': None,
            'processing_complete': False,
            'analysis_config': {
                'use_lora': True,
                'model_name': 'auto',
                'device': 'auto',
                'top_k_chunks': 5
            }
        }
        
        for key, default in defaults.items():
            if key not in st.session_state:
                st.session_state[key] = default
    
    def render_header(self):
        """Render application header."""
        st.markdown('<h1 class="main-header">📄 Research Paper Summarizer</h1>', unsafe_allow_html=True)
        st.markdown("""
        <div style="text-align: center; color: #666; margin-bottom: 2rem;">
            AI-powered analysis tool using LoRA fine-tuning for research paper summarization, 
            glossary generation, and question creation.
        </div>
        """, unsafe_allow_html=True)
    
    def render_sidebar(self):
        """Render sidebar with configuration options."""
        st.sidebar.header("⚙️ Configuration")
        
        # Model settings
        st.sidebar.subheader("Model Settings")
        
        model_options = {
            'auto': 'Auto-select (Recommended)',
            'gpt2': 'GPT-2 (CPU friendly)',
            'microsoft/phi-2': 'Phi-2 (GPU recommended)',
            'meta-llama/Llama-2-7b-chat-hf': 'LLaMA-2 7B (Large GPU required)'
        }
        
        selected_model = st.sidebar.selectbox(
            "Base Model",
            options=list(model_options.keys()),
            format_func=lambda x: model_options[x],
            index=0
        )
        
        use_lora = st.sidebar.checkbox(
            "Use LoRA Adapter",
            value=st.session_state.analysis_config['use_lora'],
            help="Enable fine-tuned LoRA adapter for better performance"
        )
        
        device_options = {
            'auto': 'Auto-detect',
            'cpu': 'CPU only',
            'cuda': 'GPU (CUDA)'
        }
        
        device = st.sidebar.selectbox(
            "Device",
            options=list(device_options.keys()),
            format_func=lambda x: device_options[x],
            index=0
        )
        
        # Processing settings
        st.sidebar.subheader("Processing Settings")
        
        top_k_chunks = st.sidebar.slider(
            "Top-K Chunks",
            min_value=3,
            max_value=10,
            value=st.session_state.analysis_config['top_k_chunks'],
            help="Number of text chunks to analyze"
        )
        
        # Update session state
        st.session_state.analysis_config.update({
            'model_name': selected_model if selected_model != 'auto' else None,
            'use_lora': use_lora,
            'device': device,
            'top_k_chunks': top_k_chunks
        })
        
        # Model info
        if st.session_state.model_loaded and st.session_state.current_model:
            st.sidebar.subheader("Current Model Info")
            model_info = st.session_state.current_model
            st.sidebar.text(f"Model: {model_info.get('model_name', 'Unknown')}")
            st.sidebar.text(f"Device: {model_info.get('device', 'Unknown')}")
            if model_info.get('lora_adapter'):
                st.sidebar.success("LoRA adapter loaded")
            else:
                st.sidebar.info("Base model only")
    
    def render_file_upload(self) -> Optional[str]:
        """Render file upload section and return temp file path if uploaded."""
        st.subheader("📁 Upload Research Paper")
        
        uploaded_file = st.file_uploader(
            "Choose a PDF file",
            type=['pdf'],
            help="Upload a research paper in PDF format"
        )
        
        if uploaded_file is not None:
            if uploaded_file != st.session_state.uploaded_file:
                # New file uploaded
                st.session_state.uploaded_file = uploaded_file
                st.session_state.analysis_results = None
                st.session_state.processing_complete = False
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                temp_path = tmp_file.name
            
            # File info
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("File Size", f"{len(uploaded_file.getvalue()) / 1024:.1f} KB")
            with col2:
                st.metric("File Name", uploaded_file.name)
            with col3:
                if st.button("🔄 Clear File"):
                    st.session_state.uploaded_file = None
                    st.session_state.analysis_results = None
                    st.session_state.processing_complete = False
                    st.experimental_rerun()
            
            return temp_path
        
        return None
    
    def render_pdf_preview(self, pdf_path: str):
        """Render PDF preview using base64 embedding."""
        st.subheader("📖 PDF Preview")
        
        try:
            with open(pdf_path, "rb") as f:
                pdf_data = f.read()
            
            base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
            pdf_display = f"""
            <div class="pdf-viewer">
                <embed
                    src="data:application/pdf;base64,{base64_pdf}"
                    width="100%"
                    height="600px"
                    type="application/pdf">
                </embed>
            </div>
            """
            
            st.markdown(pdf_display, unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"Error displaying PDF: {str(e)}")
            st.info("PDF preview unavailable, but analysis can still proceed.")
    
    def render_analysis_section(self, pdf_path: str):
        """Render analysis control section."""
        st.subheader("🤖 AI Analysis")
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.info(f"""
            **Analysis Configuration:**
            - Model: {st.session_state.analysis_config['model_name'] or 'Auto-select'}
            - LoRA: {'Enabled' if st.session_state.analysis_config['use_lora'] else 'Disabled'}
            - Device: {st.session_state.analysis_config['device'].upper()}
            - Chunks: {st.session_state.analysis_config['top_k_chunks']}
            """)
        
        with col2:
            if st.button("🚀 Analyze Paper", type="primary"):
                self.run_analysis(pdf_path)
    
    def run_analysis(self, pdf_path: str):
        """Run the paper analysis pipeline."""
        with st.spinner("🔄 Analyzing paper... This may take a few minutes."):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Step 1: Process PDF
                status_text.text("📄 Processing PDF...")
                progress_bar.progress(20)
                
                # Step 2: Load model
                status_text.text("🤖 Loading AI model...")
                progress_bar.progress(40)
                
                # Step 3: Run analysis
                status_text.text("🧠 Generating analysis...")
                progress_bar.progress(60)
                
                results = run_analysis(
                    input_path=pdf_path,
                    task_type='all',
                    model_name=st.session_state.analysis_config['model_name'],
                    device=st.session_state.analysis_config['device'],
                    use_lora=st.session_state.analysis_config['use_lora'],
                    top_k=st.session_state.analysis_config['top_k_chunks']
                )
                
                progress_bar.progress(100)
                status_text.text("✅ Analysis complete!")
                
                # Store results
                st.session_state.analysis_results = results
                st.session_state.current_model = results.get('model_info', {})
                st.session_state.model_loaded = True
                st.session_state.processing_complete = True
                
                time.sleep(1)
                status_text.empty()
                progress_bar.empty()
                
                st.success("🎉 Analysis completed successfully!")
                
            except Exception as e:
                progress_bar.empty()
                status_text.empty()
                
                st.error(f"❌ Analysis failed: {str(e)}")
                
                with st.expander("🔧 Debug Information"):
                    st.code(traceback.format_exc())
    
    def render_results(self):
        """Render analysis results."""
        if not st.session_state.analysis_results:
            st.info("👆 Upload a PDF and click 'Analyze Paper' to see results here.")
            return
        
        results = st.session_state.analysis_results
        
        st.subheader("📊 Analysis Results")
        
        # Results tabs
        tab1, tab2, tab3, tab4 = st.tabs(["📝 Summary", "📚 Glossary", "❓ Questions", "📈 Metrics"])
        
        with tab1:
            self.render_summary_tab(results)
        
        with tab2:
            self.render_glossary_tab(results)
        
        with tab3:
            self.render_questions_tab(results)
        
        with tab4:
            self.render_metrics_tab(results)
        
        # Download section
        st.subheader("💾 Download Results")
        self.render_download_section(results)
    
    def render_summary_tab(self, results: Dict[str, Any]):
        """Render summary results tab."""
        summary_result = self.get_task_result(results, 'summary')
        
        if summary_result:
            st.markdown('<div class="task-header">5-Point Summary</div>', unsafe_allow_html=True)
            
            output = summary_result['output']
            
            # Display summary points
            summary_lines = output.split('\n')
            for line in summary_lines:
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('•')):
                    st.markdown(f"**{line}**")
                elif line:
                    st.markdown(line)
            
            # Validation info
            validation = summary_result.get('validation', {})
            if validation:
                self.render_validation_info(validation, 'summary')
        else:
            st.warning("No summary results available.")
    
    def render_glossary_tab(self, results: Dict[str, Any]):
        """Render glossary results tab."""
        glossary_result = self.get_task_result(results, 'glossary')
        
        if glossary_result:
            st.markdown('<div class="task-header">Technical Glossary</div>', unsafe_allow_html=True)
            
            output = glossary_result['output']
            
            # Parse and display glossary terms
            terms = []
            for line in output.split('\n'):
                line = line.strip()
                if ':' in line:
                    term, definition = line.split(':', 1)
                    terms.append((term.strip(), definition.strip()))
            
            if terms:
                for term, definition in terms:
                    with st.container():
                        st.markdown(f"**{term}**")
                        st.markdown(f"<div style='margin-left: 1rem; color: #555;'>{definition}</div>", 
                                  unsafe_allow_html=True)
                        st.markdown("---")
            else:
                st.text(output)
            
            # Validation info
            validation = glossary_result.get('validation', {})
            if validation:
                self.render_validation_info(validation, 'glossary')
        else:
            st.warning("No glossary results available.")
    
    def render_questions_tab(self, results: Dict[str, Any]):
        """Render questions results tab."""
        questions_result = self.get_task_result(results, 'questions')
        
        if questions_result:
            st.markdown('<div class="task-header">Exam Questions</div>', unsafe_allow_html=True)
            
            output = questions_result['output']
            
            # Parse and display questions
            questions = output.split('\n\n') if '\n\n' in output else output.split('\n')
            
            for i, question in enumerate(questions, 1):
                question = question.strip()
                if question:
                    with st.container():
                        # Check if it's a multiple choice question
                        if '(A)' in question or '(B)' in question:
                            st.markdown(f"**Question {i} (Multiple Choice):**")
                            
                            # Highlight correct answer if present
                            if '[Answer:' in question:
                                answer_match = question.split('[Answer:')
                                question_part = answer_match[0].strip()
                                answer_part = answer_match[1].strip(' ]')
                                
                                st.markdown(question_part)
                                st.success(f"**Correct Answer: {answer_part}**")
                            else:
                                st.markdown(question)
                        else:
                            st.markdown(f"**Question {i}:**")
                            st.markdown(question)
                        
                        st.markdown("---")
            
            # Validation info
            validation = questions_result.get('validation', {})
            if validation:
                self.render_validation_info(validation, 'questions')
        else:
            st.warning("No questions results available.")
    
    def render_metrics_tab(self, results: Dict[str, Any]):
        """Render metrics and performance tab."""
        st.markdown('<div class="task-header">Performance Metrics</div>', unsafe_allow_html=True)
        
        # Model info
        model_info = results.get('model_info', {})
        if model_info:
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Model", model_info.get('model_name', 'Unknown'))
                
            with col2:
                st.metric("Device", model_info.get('device', 'Unknown'))
                
            with col3:
                lora_status = "✅ Enabled" if model_info.get('lora_adapter') else "❌ Disabled"
                st.metric("LoRA Adapter", lora_status)
        
        # Processing stats
        stats = results.get('statistics', {})
        if stats:
            st.subheader("📄 Document Statistics")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Pages", stats.get('total_pages', 0))
            
            with col2:
                st.metric("Characters", f"{stats.get('total_chars', 0):,}")
            
            with col3:
                st.metric("Total Chunks", stats.get('total_chunks', 0))
            
            with col4:
                st.metric("Selected Chunks", stats.get('selected_chunks', 0))
        
        # Timing info
        st.subheader("⏱️ Processing Time")
        
        total_time = results.get('total_time', 0)
        st.metric("Total Analysis Time", f"{total_time:.1f}s")
        
        # Task-specific metrics
        task_results = results.get('task_results', [])
        if task_results:
            st.subheader("🎯 Task Performance")
            
            for task_result in task_results:
                task_type = task_result['task_type'].title()
                metadata = task_result.get('metadata', {})
                
                exec_time = metadata.get('total_execution_time', 0)
                tokens_per_sec = metadata.get('tokens_per_second', 0)
                
                with st.container():
                    st.markdown(f"**{task_type}:**")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.text(f"Execution time: {exec_time:.1f}s")
                    with col2:
                        st.text(f"Tokens/sec: {tokens_per_sec:.1f}")
                    
                    st.markdown("---")
    
    def render_validation_info(self, validation: Dict[str, Any], task_type: str):
        """Render validation information."""
        if not validation.get('valid', True):
            with st.expander("⚠️ Validation Issues"):
                for issue in validation.get('issues', []):
                    st.warning(f"• {issue}")
        
        score = validation.get('score', 1.0)
        if score < 1.0:
            st.info(f"Quality Score: {score:.1%}")
    
    def render_download_section(self, results: Dict[str, Any]):
        """Render download options."""
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # JSON download
            json_data = json.dumps(results, indent=2, ensure_ascii=False)
            st.download_button(
                label="📄 Download JSON",
                data=json_data,
                file_name=f"analysis_results_{int(time.time())}.json",
                mime="application/json"
            )
        
        with col2:
            # TXT download
            txt_content = self.create_text_report(results)
            st.download_button(
                label="📝 Download TXT",
                data=txt_content,
                file_name=f"analysis_report_{int(time.time())}.txt",
                mime="text/plain"
            )
        
        with col3:
            # Word document download (if available)
            if DOCX_AVAILABLE:
                try:
                    docx_content = self.create_word_report(results)
                    st.download_button(
                        label="📊 Download DOCX",
                        data=docx_content,
                        file_name=f"analysis_report_{int(time.time())}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"DOCX generation failed: {e}")
            else:
                st.info("Install python-docx for Word export")
    
    def get_task_result(self, results: Dict[str, Any], task_type: str) -> Optional[Dict[str, Any]]:
        """Get result for specific task type."""
        task_results = results.get('task_results', [])
        for result in task_results:
            if result.get('task_type') == task_type:
                return result
        return None
    
    def create_text_report(self, results: Dict[str, Any]) -> str:
        """Create text report from results."""
        report = []
        
        # Header
        report.append("RESEARCH PAPER ANALYSIS REPORT")
        report.append("=" * 50)
        report.append("")
        
        # Metadata
        metadata = results.get('paper_metadata', {})
        if metadata.get('title'):
            report.append(f"Title: {metadata['title']}")
        report.append(f"Analysis Date: {results.get('timestamp', 'Unknown')}")
        report.append("")
        
        # Results for each task
        task_results = results.get('task_results', [])
        for task_result in task_results:
            task_type = task_result['task_type'].upper()
            output = task_result['output']
            
            report.append(f"{task_type}")
            report.append("-" * len(task_type))
            report.append(output)
            report.append("")
        
        # Model info
        model_info = results.get('model_info', {})
        if model_info:
            report.append("MODEL INFORMATION")
            report.append("-" * 18)
            report.append(f"Model: {model_info.get('model_name', 'Unknown')}")
            report.append(f"Device: {model_info.get('device', 'Unknown')}")
            report.append(f"LoRA Adapter: {'Yes' if model_info.get('lora_adapter') else 'No'}")
            report.append("")
        
        return "\n".join(report)
    
    def create_word_report(self, results: Dict[str, Any]) -> bytes:
        """Create Word document report."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Research Paper Analysis Report', 0)
        title.alignment = 1  # Center alignment
        
        # Metadata
        metadata = results.get('paper_metadata', {})
        if metadata.get('title'):
            doc.add_heading(f"Paper: {metadata['title']}", level=1)
        
        doc.add_paragraph(f"Analysis Date: {results.get('timestamp', 'Unknown')}")
        doc.add_paragraph("")
        
        # Task results
        task_results = results.get('task_results', [])
        for task_result in task_results:
            task_type = task_result['task_type'].title()
            output = task_result['output']
            
            doc.add_heading(task_type, level=1)
            
            # Format output based on task type
            if task_result['task_type'] == 'summary':
                # Format as numbered list
                lines = output.split('\n')
                for line in lines:
                    if line.strip():
                        doc.add_paragraph(line.strip())
            
            elif task_result['task_type'] == 'glossary':
                # Format as definition list
                for line in output.split('\n'):
                    if ':' in line:
                        term, definition = line.split(':', 1)
                        p = doc.add_paragraph()
                        p.add_run(term.strip()).bold = True
                        p.add_run(f": {definition.strip()}")
            
            else:
                # Default formatting
                doc.add_paragraph(output)
            
            doc.add_paragraph("")
        
        # Model information
        model_info = results.get('model_info', {})
        if model_info:
            doc.add_heading('Model Information', level=1)
            doc.add_paragraph(f"Model: {model_info.get('model_name', 'Unknown')}")
            doc.add_paragraph(f"Device: {model_info.get('device', 'Unknown')}")
            doc.add_paragraph(f"LoRA Adapter: {'Enabled' if model_info.get('lora_adapter') else 'Disabled'}")
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue()
    
    def run(self):
        """Main application run method."""
        self.render_header()
        self.render_sidebar()
        
        # Main content in two columns
        col1, col2 = st.columns([1, 1])
        
        with col1:
            # Left side: Upload and PDF preview
            pdf_path = self.render_file_upload()
            
            if pdf_path:
                self.render_pdf_preview(pdf_path)
                self.render_analysis_section(pdf_path)
        
        with col2:
            # Right side: Results
            self.render_results()


def main():
    """Main function to run the Streamlit app."""
    app = StreamlitApp()
    app.run()


if __name__ == "__main__":
    main()